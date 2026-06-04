import os
import pypandoc
from docx import Document
from docx.shared import Inches

def setup_pandoc():
    try:
        pypandoc.get_pandoc_version()
        print("Pandoc is already installed.")
    except OSError:
        print("Downloading pandoc...")
        pypandoc.download_pandoc()

def convert_md_to_docx():
    docs_dir = 'docs'
    md_files = [f for f in os.listdir(docs_dir) if f.endswith('.md')]
    
    for md_file in md_files:
        md_path = os.path.join(docs_dir, md_file)
        docx_path = os.path.join(docs_dir, md_file.replace('.md', '.docx'))
        print(f"Converting {md_file} to {docx_path}...")
        
        # Pypandoc conversion
        try:
            # We don't embed images directly via pandoc easily if paths are relative sometimes, but we can try
            pypandoc.convert_file(md_path, 'docx', outputfile=docx_path, extra_args=['--reference-doc=custom-reference.docx'] if os.path.exists('custom-reference.docx') else [])
        except Exception as e:
            print(f"Error converting {md_file}: {e}")

def build_manuel():
    print("Building Manuel Utilisateur with images...")
    doc = Document()
    doc.add_heading('Manuel d\'Utilisation - UniGames', 0)
    
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('Bienvenue sur la plateforme UniGames.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214135_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214135_localhost.jpeg', width=Inches(6.0))
        
    doc.add_heading('2. Tableau de Bord', level=1)
    doc.add_paragraph('Le tableau de bord permet de voir les statistiques de l\'édition.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214158_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214158_localhost.jpeg', width=Inches(6.0))
        
    doc.add_heading('3. Éditions et Disciplines', level=1)
    doc.add_paragraph('Gestion globale des événements et des sports.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214429_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214429_localhost.jpeg', width=Inches(6.0))
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214443_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214443_localhost.jpeg', width=Inches(6.0))
        
    doc.add_heading('4. Facultés et Équipes', level=1)
    doc.add_paragraph('Création des universités et des équipes.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214537_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214537_localhost.jpeg', width=Inches(6.0))
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_21459_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_21459_localhost.jpeg', width=Inches(6.0))

    doc.add_heading('5. Joueurs', level=1)
    doc.add_paragraph('Gestion des effectifs.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214710_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214710_localhost.jpeg', width=Inches(6.0))

    doc.add_heading('6. Programmation des Matchs', level=1)
    doc.add_paragraph('Suivi et scores en direct.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214920_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214920_localhost.jpeg', width=Inches(6.0))
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214952_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214952_localhost.jpeg', width=Inches(6.0))

    doc.add_heading('7. Classements', level=1)
    doc.add_paragraph('Génération automatique des classements.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_215027_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_215027_localhost.jpeg', width=Inches(6.0))
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_21560_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_21560_localhost.jpeg', width=Inches(6.0))

    doc.save('docs/05_manuel_utilisateur.docx')
    print("Manuel utilisateur generated.")

if __name__ == '__main__':
    setup_pandoc()
    convert_md_to_docx()
    # Override Manuel Utilisateur with the picture-embedded version
    build_manuel()
    print("All tasks completed.")
