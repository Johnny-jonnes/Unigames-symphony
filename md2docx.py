import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = 'docs'
IMAGES_DIR = os.path.join(DOCS_DIR, 'images')

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('![') or '![' in line:
            # Try to extract image path
            m = re.search(r'!\[.*?\]\((.*?)\)', line)
            if m:
                img_path = m.group(1)
                img_path = img_path.replace('%20', ' ')
                
                # Check absolute paths
                if img_path.startswith('/C:'):
                    img_path = img_path[1:] # remove leading slash
                elif not os.path.isabs(img_path):
                    img_path = os.path.join(DOCS_DIR, img_path)
                
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                    except Exception as e:
                        print(f"Erreur d'ajout de l'image {img_path} : {e}")
                else:
                    print(f"Image introuvable : {img_path}")
        elif line.startswith('>') or line.startswith('```') or line.startswith('|'):
            # Simple format for blocks and tables
            p = doc.add_paragraph(line)
            p.style = 'Quote'
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Converti: {docx_path}")

def main():
    if not os.path.exists(DOCS_DIR):
        print("Dossier docs/ introuvable.")
        return
        
    for f in os.listdir(DOCS_DIR):
        if f.endswith('.md'):
            md_path = os.path.join(DOCS_DIR, f)
            docx_path = os.path.join(DOCS_DIR, f.replace('.md', '.docx'))
            md_to_docx(md_path, docx_path)

if __name__ == '__main__':
    main()
