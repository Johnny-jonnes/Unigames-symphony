import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('Manuel d\'Utilisation - UniGames', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Date : 4 Juin 2026')
    doc.add_paragraph('Projet : UniGames')
    
    doc.add_heading('1. Introduction et Prérequis', level=1)
    doc.add_paragraph('La plateforme UniGames est accessible via n\'importe quel navigateur web moderne. Connectez-vous avec vos identifiants fournis par l\'administrateur.')
    
    doc.add_heading('2. Tableau de Bord (Dashboard)', level=1)
    doc.add_paragraph('Dès votre connexion, vous arrivez sur le Tableau de bord qui affiche les statistiques globales.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214158_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214158_localhost.jpeg', width=Inches(6.0))
        
    doc.add_heading('3. Gestion des Éditions et Disciplines', level=1)
    doc.add_paragraph('L\'administration permet de configurer les éditions annuelles et leurs disciplines respectives.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214429_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214429_localhost.jpeg', width=Inches(6.0))
        
    doc.add_heading('4. Programmation et Saisie de Matchs', level=1)
    doc.add_paragraph('Planifiez les matchs et saisissez les scores et buteurs en temps réel.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_214920_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_214920_localhost.jpeg', width=Inches(6.0))

    doc.add_heading('5. Classements', level=1)
    doc.add_paragraph('Les classements sont calculés dynamiquement selon l\'édition et la discipline sélectionnées.')
    if os.path.exists('docs/images/Capture d’écran_4-6-2026_215027_localhost.jpeg'):
        doc.add_picture('docs/images/Capture d’écran_4-6-2026_215027_localhost.jpeg', width=Inches(6.0))

    doc.save('docs/Manuel_Utilisateur.docx')

def create_portfolio():
    doc = Document()
    
    title = doc.add_heading('Portfolio d\'Interface - UniGames', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Ce document présente l\'interface finale de l\'application web UniGames (Thème Sombre, Responsive).')
    
    images = [f for f in os.listdir('docs/images') if f.endswith(('.jpeg', '.jpg', '.png'))]
    images.sort()
    
    for img in images:
        doc.add_heading(img, level=2)
        try:
            doc.add_picture(os.path.join('docs/images', img), width=Inches(6.0))
        except Exception as e:
            print(f"Error adding {img}: {e}")
            
    doc.save('docs/Maquettes_Graphiques.docx')

if __name__ == '__main__':
    create_manual()
    create_portfolio()
    print("Documents Word generes avec succes !")
